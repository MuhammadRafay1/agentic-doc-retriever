"""
Utility module for AI Research Assistant
Handles document loading, text processing, embedding generation, and vector store management.
"""

import os
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import logging

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS, Chroma
from langchain.docstore.document import Document

# Document loading
import PyPDF2
from docx import Document as DocxDocument

from app.config import (
    CHUNK_SIZE, 
    CHUNK_OVERLAP, 
    SUPPORTED_FORMATS,
    VECTOR_STORE_DIR,
    EMBEDDING_MODELS
)

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentLoader:
    """Handles loading documents from various file formats."""
    
    @staticmethod
    def load_txt(file_path: Path) -> str:
        """Load text from .txt file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @staticmethod
    def load_pdf(file_path: Path) -> str:
        """Load text from .pdf file."""
        text = ""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
        return text
    
    @staticmethod
    def load_docx(file_path: Path) -> str:
        """Load text from .docx file."""
        try:
            doc = DocxDocument(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            return ""
    
    @staticmethod
    def load_md(file_path: Path) -> str:
        """Load text from .md file."""
        return DocumentLoader.load_txt(file_path)
    
    @classmethod
    def load_document(cls, file_path: Path) -> Optional[Document]:
        """
        Load a document from file and return as LangChain Document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            LangChain Document object or None if loading fails
        """
        suffix = file_path.suffix.lower()
        
        loaders = {
            '.txt': cls.load_txt,
            '.pdf': cls.load_pdf,
            '.docx': cls.load_docx,
            '.md': cls.load_md
        }
        
        if suffix not in loaders:
            logger.warning(f"Unsupported format: {suffix}")
            return None
        
        text = loaders[suffix](file_path)
        if text.strip():
            return Document(
                page_content=text,
                metadata={"source": str(file_path), "filename": file_path.name}
            )
        return None
    
    @classmethod
    def load_documents_from_directory(cls, directory: Path) -> Tuple[List[Document], Dict]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory: Path to directory containing documents
            
        Returns:
            Tuple of (list of Documents, statistics dictionary)
        """
        documents = []
        stats = {
            'total_files': 0,
            'loaded_files': 0,
            'failed_files': 0,
            'total_size_bytes': 0,
            'file_types': {}
        }
        
        for file_path in Path(directory).rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_FORMATS:
                stats['total_files'] += 1
                stats['total_size_bytes'] += file_path.stat().st_size
                
                # Track file type
                ext = file_path.suffix.lower()
                stats['file_types'][ext] = stats['file_types'].get(ext, 0) + 1
                
                doc = cls.load_document(file_path)
                if doc:
                    documents.append(doc)
                    stats['loaded_files'] += 1
                else:
                    stats['failed_files'] += 1
        
        logger.info(f"Loaded {stats['loaded_files']}/{stats['total_files']} documents")
        return documents, stats


class TextProcessor:
    """Handles text chunking and processing."""
    
    def __init__(self, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP):
        """
        Initialize text processor.
        
        Args:
            chunk_size: Size of text chunks in characters
            chunk_overlap: Overlap between chunks in characters
        """
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        """
        Split documents into smaller chunks.
        
        Args:
            documents: List of LangChain Documents
            
        Returns:
            List of chunked Documents
        """
        chunks = self.text_splitter.split_documents(documents)
        logger.info(f"Split {len(documents)} documents into {len(chunks)} chunks")
        return chunks


class EmbeddingEngine:
    """Handles embedding generation using Hugging Face models."""
    
    def __init__(self, model_key: str):
        """
        Initialize embedding engine with specified model.
        
        Args:
            model_key: Key from EMBEDDING_MODELS config
        """
        if model_key not in EMBEDDING_MODELS:
            raise ValueError(f"Unknown model: {model_key}")
        
        self.model_key = model_key
        self.model_name = EMBEDDING_MODELS[model_key]["name"]
        
        logger.info(f"Loading embedding model: {self.model_name}")
        self.embeddings = HuggingFaceEmbeddings(
            model_name=self.model_name,
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Embed a list of documents."""
        return self.embeddings.embed_documents(texts)
    
    def embed_query(self, text: str) -> List[float]:
        """Embed a single query."""
        return self.embeddings.embed_query(text)


class VectorStoreManager:
    """Manages vector store creation, saving, and loading."""
    
    def __init__(self, store_type: str, embedding_engine: EmbeddingEngine):
        """
        Initialize vector store manager.
        
        Args:
            store_type: Type of vector store ('FAISS' or 'ChromaDB')
            embedding_engine: Initialized EmbeddingEngine instance
        """
        self.store_type = store_type
        self.embedding_engine = embedding_engine
        self.vector_store = None
    
    def create_vector_store(self, documents: List[Document]) -> None:
        """
        Create vector store from documents.
        
        Args:
            documents: List of LangChain Document chunks
        """
        logger.info(f"Creating {self.store_type} vector store from {len(documents)} chunks")
        
        if self.store_type == "FAISS":
            self.vector_store = FAISS.from_documents(
                documents=documents,
                embedding=self.embedding_engine.embeddings
            )
        elif self.store_type == "ChromaDB":
            self.vector_store = Chroma.from_documents(
                documents=documents,
                embedding=self.embedding_engine.embeddings,
                persist_directory=str(VECTOR_STORE_DIR / "chroma_db")
            )
        else:
            raise ValueError(f"Unsupported vector store: {self.store_type}")
        
        logger.info(f"{self.store_type} vector store created successfully")
    
    def save_vector_store(self, name: str) -> None:
        """
        Save vector store to disk.
        
        Args:
            name: Name for the saved store
        """
        if self.vector_store is None:
            raise ValueError("No vector store to save")
        
        save_path = VECTOR_STORE_DIR / f"{name}_{self.store_type.lower()}"
        
        if self.store_type == "FAISS":
            self.vector_store.save_local(str(save_path))
            logger.info(f"FAISS store saved to {save_path}")
        elif self.store_type == "ChromaDB":
            # ChromaDB persists automatically if persist_directory is set
            logger.info(f"ChromaDB persisted to {VECTOR_STORE_DIR / 'chroma_db'}")
    
    def load_vector_store(self, name: str) -> bool:
        """
        Load vector store from disk.
        
        Args:
            name: Name of the saved store
            
        Returns:
            True if loaded successfully, False otherwise
        """
        load_path = VECTOR_STORE_DIR / f"{name}_{self.store_type.lower()}"
        
        try:
            if self.store_type == "FAISS":
                self.vector_store = FAISS.load_local(
                    str(load_path),
                    self.embedding_engine.embeddings,
                    allow_dangerous_deserialization=True
                )
                logger.info(f"FAISS store loaded from {load_path}")
                return True
            elif self.store_type == "ChromaDB":
                self.vector_store = Chroma(
                    persist_directory=str(VECTOR_STORE_DIR / "chroma_db"),
                    embedding_function=self.embedding_engine.embeddings
                )
                logger.info(f"ChromaDB loaded from {VECTOR_STORE_DIR / 'chroma_db'}")
                return True
        except Exception as e:
            logger.error(f"Error loading vector store: {e}")
            return False
        
        return False
    
    def similarity_search(self, query: str, k: int = 5) -> List[Tuple[Document, float]]:
        """
        Perform similarity search on vector store.
        
        Args:
            query: Search query
            k: Number of top results to return
            
        Returns:
            List of (Document, similarity_score) tuples
        """
        if self.vector_store is None:
            raise ValueError("No vector store loaded")
        
        # Perform search with scores
        results = self.vector_store.similarity_search_with_score(query, k=k)
        
        logger.info(f"Found {len(results)} results for query: {query[:50]}...")
        return results


# Convenience function for quick setup
def create_semantic_search_system(
    data_directory: Path,
    embedding_model: str,
    vector_store_type: str
) -> Tuple[VectorStoreManager, Dict]:
    """
    Create complete semantic search system from directory.
    
    Args:
        data_directory: Path to directory with documents
        embedding_model: Embedding model key
        vector_store_type: Type of vector store
        
    Returns:
        Tuple of (VectorStoreManager, statistics)
    """
    # Load documents
    documents, stats = DocumentLoader.load_documents_from_directory(data_directory)
    
    if not documents:
        raise ValueError("No documents loaded from directory")
    
    # Process documents
    processor = TextProcessor()
    chunks = processor.split_documents(documents)
    
    # Create embeddings
    embedding_engine = EmbeddingEngine(embedding_model)
    
    # Create and populate vector store
    vector_manager = VectorStoreManager(vector_store_type, embedding_engine)
    vector_manager.create_vector_store(chunks)
    
    stats['total_chunks'] = len(chunks)
    return vector_manager, stats
